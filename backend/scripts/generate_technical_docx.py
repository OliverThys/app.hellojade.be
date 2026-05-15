#!/usr/bin/env python3
"""
Génère docs/HelloJADE_documentation_technique.docx à partir des Markdown dans docs/technical/.

- Page de garde + logo HelloJADE
- Table des matières (titres niveau 1 extraits des fichiers)
- Corps : fusion des fichiers 00-…12 triés + README optionnel
- Blocs ```mermaid → image via https://kroki.io (désactiver : HELLOJADE_DOCX_KROKI=0)

Usage :
  pip install python-docx
  python backend/scripts/generate_technical_docx.py
"""
from __future__ import annotations

import os
import re
import sys
import tempfile
import urllib.error
import urllib.request
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, List

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("Installez python-docx : pip install python-docx", file=sys.stderr)
    sys.exit(1)


REPO_ROOT = Path(__file__).resolve().parents[2]
TECH_DOCS = REPO_ROOT / "docs" / "technical"
OUTPUT_PATH = REPO_ROOT / "docs" / "HelloJADE_documentation_technique.docx"
KROKI_URL = os.environ.get("HELLOJADE_KROKI_URL", "https://kroki.io/mermaid/png")
USE_KROKI = os.environ.get("HELLOJADE_DOCX_KROKI", "1").lower() not in ("0", "false", "no")

LOGO_CANDIDATES = [
    REPO_ROOT / "frontend" / "src" / "assets" / "hellojade_logo.png",
    REPO_ROOT / "backend" / "app" / "assets" / "hellojade_logo.png",
    REPO_ROOT / "hellojade_logo.png",
]


def _find_logo() -> Path | None:
    for p in LOGO_CANDIDATES:
        if p.is_file():
            return p
    return None


def _set_cell_shading(cell: Any, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _style_doc(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)


def _add_cover(document: Document, logo_path: Path | None) -> None:
    for _ in range(2):
        document.add_paragraph()

    if logo_path:
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Cm(4.2))
    else:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[Logo HelloJADE — voir frontend/src/assets/hellojade_logo.png]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_paragraph()
    t = document.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("Documentation technique")
    rt.bold = True
    rt.font.size = Pt(28)
    rt.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    st = document.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = st.add_run("HelloJADE — Suivi post-hospitalisation assisté par IA")
    rs.font.size = Pt(14)
    rs.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    document.add_paragraph()
    ep = document.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re = ep.add_run("Document destiné au partenariat technique\nCentre Hospitalier Epi CURA (Epicura)")
    re.font.size = Pt(12)

    document.add_paragraph()
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"Version document : 1.0\nDate de génération : {date.today().strftime('%d/%m/%Y')}\n"
        "Sources : dossier docs/technical/ (*.md)\n"
        "Classification : usage interne / partenaires autorisés"
    )
    rm.font.size = Pt(10)
    rm.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    document.add_page_break()


def _kroki_diagram_png(diagram_type: str, source: str) -> bytes | None:
    """POST diagramme vers Kroki (ex. type=mermaid)."""
    if not USE_KROKI or not source.strip():
        return None
    url = KROKI_URL
    if diagram_type != "mermaid" and "/mermaid/" in url:
        url = url.replace("/mermaid/", f"/{diagram_type}/")
    try:
        req = urllib.request.Request(
            url,
            data=source.encode("utf-8"),
            method="POST",
            headers={"Content-Type": "text/plain; charset=utf-8"},
        )
        with urllib.request.urlopen(req, timeout=45) as resp:
            return resp.read()
    except (urllib.error.URLError, urllib.error.HTTPError, TimeoutError, OSError):
        return None


def _strip_front_matter(lines: List[str]) -> List[str]:
    if not lines or lines[0].strip() != "---":
        return lines
    i = 1
    while i < len(lines) and lines[i].strip() != "---":
        i += 1
    if i < len(lines):
        return lines[i + 1 :]
    return lines


@dataclass
class MdBlock:
    kind: str
    data: Any


def _parse_markdown(text: str) -> List[MdBlock]:
    lines = _strip_front_matter(text.splitlines())
    blocks: List[MdBlock] = []
    i = 0
    n = len(lines)

    def is_table_row(s: str) -> bool:
        s = s.strip()
        return s.startswith("|") and s.endswith("|") and s.count("|") >= 2

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip() or "text"
            i += 1
            buf: List[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < n:
                i += 1
            blocks.append(MdBlock("fence", (lang, "\n".join(buf))))
            continue

        if is_table_row(line):
            rows: List[List[str]] = []
            while i < n and is_table_row(lines[i]):
                cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            if rows:
                blocks.append(MdBlock("table", rows))
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            blocks.append(MdBlock("h", (level, title)))
            i += 1
            continue

        if re.match(r"^[-*]\s+", line.strip()):
            items: List[str] = []
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(MdBlock("ul", items))
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            items = []
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(MdBlock("ol", items))
            continue

        if line.strip() in ("---", "***", "___") and len(line.strip()) >= 3:
            blocks.append(MdBlock("hr", None))
            i += 1
            continue

        if line.strip().startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append(MdBlock("quote", "\n".join(quote_lines)))
            continue

        para: List[str] = []
        while i < n:
            ln = lines[i].rstrip()
            if not ln.strip():
                break
            if ln.strip().startswith("```"):
                break
            if ln.strip().startswith("#"):
                break
            if is_table_row(ln):
                break
            if re.match(r"^[-*]\s+", ln.strip()):
                break
            if re.match(r"^\d+\.\s+", ln.strip()):
                break
            if ln.strip() in ("---", "***", "___") and len(ln.strip()) >= 3:
                break
            if ln.strip().startswith(">"):
                break
            para.append(ln)
            i += 1
        if para:
            blocks.append(MdBlock("p", "\n".join(para)))
        else:
            i += 1
        continue

    return blocks


def _add_runs_with_bold(paragraph: Any, text: str) -> None:
    """Découpe **gras** et `code` simple."""
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            continue
        sub = re.split(r"(\*\*[^*]+\*\*)", part)
        for s in sub:
            if s.startswith("**") and s.endswith("**"):
                r = paragraph.add_run(s[2:-2])
                r.bold = True
            elif s:
                paragraph.add_run(s)


def _emit_blocks(document: Document, blocks: List[MdBlock]) -> None:
    for b in blocks:
        if b.kind == "h":
            level, title = b.data
            doc_level = min(max(level, 1), 9)
            document.add_heading(title, level=doc_level)
        elif b.kind == "p":
            p = document.add_paragraph()
            for i, line in enumerate(b.data.split("\n")):
                if i:
                    p.add_run().add_break()
                _add_runs_with_bold(p, line)
        elif b.kind == "ul":
            for item in b.data:
                p = document.add_paragraph(style="List Bullet")
                _add_runs_with_bold(p, item)
        elif b.kind == "ol":
            for item in b.data:
                p = document.add_paragraph(style="List Number")
                _add_runs_with_bold(p, item)
        elif b.kind == "quote":
            p = document.add_paragraph(style="Intense Quote")
            _add_runs_with_bold(p, b.data)
        elif b.kind == "hr":
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 42)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif b.kind == "table":
            rows: List[List[str]] = b.data
            if not rows:
                continue
            ncol = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=ncol)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(ncol):
                    cell = table.rows[ri].cells[ci]
                    cell.text = row[ci] if ci < len(row) else ""
                    if ri == 0:
                        for pr in cell.paragraphs:
                            for r in pr.runs:
                                r.bold = True
                        _set_cell_shading(cell, "E8F4FC")
            document.add_paragraph()
        elif b.kind == "fence":
            lang, body = b.data
            body = body.rstrip("\n")
            if lang == "mermaid" and body.strip():
                document.add_paragraph()
                cap = document.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run("Figure — diagramme")
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                png = _kroki_diagram_png("mermaid", body)
                if png:
                    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    try:
                        tmp.write(png)
                        tmp.flush()
                        tmp.close()
                        picp = document.add_paragraph()
                        picp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pr = picp.add_run()
                        pr.add_picture(tmp.name, width=Cm(15.5))
                    finally:
                        Path(tmp.name).unlink(missing_ok=True)
                else:
                    note = document.add_paragraph()
                    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    nr = note.add_run(
                        "(Image Mermaid non générée — réseau indisponible ou HELLOJADE_DOCX_KROKI=0. "
                        "Collez le source dans https://kroki.io ou regénérez avec accès Internet.)"
                    )
                    nr.font.size = Pt(9)
                    nr.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
                srcp = document.add_paragraph()
                sr = srcp.add_run(body)
                sr.font.name = "Consolas"
                sr.font.size = Pt(8)
                srcp.paragraph_format.left_indent = Cm(0.4)
                document.add_paragraph()
            else:
                document.add_paragraph()
                lab = document.add_paragraph()
                lr = lab.add_run(f"Listing ({lang})")
                lr.bold = True
                lr.font.size = Pt(9)
                cp = document.add_paragraph()
                cr = cp.add_run(body if body else " ")
                cr.font.name = "Consolas"
                cr.font.size = Pt(9)
                cp.paragraph_format.left_indent = Cm(0.5)
                document.add_paragraph()


def _first_h1(text: str) -> str | None:
    for line in text.splitlines():
        m = re.match(r"^#\s+(.+)$", line.strip())
        if m:
            return m.group(1).strip()
    return None


def _collect_md_files() -> List[Path]:
    readme = TECH_DOCS / "README.md"
    numbered = sorted(TECH_DOCS.glob("[0-9][0-9]-*.md"))
    out: List[Path] = []
    if readme.is_file():
        out.append(readme)
    out.extend(numbered)
    return out


def _add_toc_from_files(document: Document, files: List[Path]) -> None:
    document.add_heading("Table des matières", level=1)
    document.add_paragraph(
        "Les chapitres suivants reprennent les fichiers Markdown du dépôt (dossier docs/technical/). "
        "Sous Word : après ouverture du document, onglet Références → Table des matières → modèle automatique "
        "pour insérer une TDM à numéros de page si besoin d’une version imprimable."
    )
    for p in files:
        if p.name == "README.md":
            document.add_paragraph("Guide de lecture — README", style="List Number")
            continue
        h = _first_h1(p.read_text(encoding="utf-8"))
        label = h or p.stem.replace("-", " ").title()
        document.add_paragraph(label, style="List Number")
    document.add_page_break()


def _add_footer(document: Document) -> None:
    section = document.sections[-1]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fr = fp.add_run(
        f"HelloJADE — Documentation technique — généré le {date.today().isoformat()} — "
        "Sources : docs/technical/*.md — Usage interne / partenaires autorisés"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> int:
    if not TECH_DOCS.is_dir():
        print(f"Dossier manquant : {TECH_DOCS}", file=sys.stderr)
        return 1

    files = _collect_md_files()
    if not files:
        print(f"Aucun fichier .md dans {TECH_DOCS}", file=sys.stderr)
        return 1

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    logo = _find_logo()

    document = Document()
    _style_doc(document)
    _add_cover(document, logo)
    _add_toc_from_files(document, files)

    for idx, path in enumerate(files):
        text = path.read_text(encoding="utf-8")
        document.add_heading(f"Fichier source : {path.name}", level=2)
        blocks = _parse_markdown(text)
        _emit_blocks(document, blocks)
        if idx < len(files) - 1:
            document.add_page_break()

    # Retirer la dernière page break inutile : remplacer par paragraphe vide si possible
    # (python-docx ne permet pas facilement de pop le dernier élément) — acceptable

    _add_footer(document)
    document.save(str(OUTPUT_PATH))
    print(f"Écrit : {OUTPUT_PATH} ({len(files)} fichiers Markdown)")
    if not USE_KROKI:
        print("Note : HELLOJADE_DOCX_KROKI=0 — diagrammes Mermaid en texte uniquement.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
